# -*- coding: utf-8 -*-

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "okb_mvp_weekly_checklist_ru.md"
OUT_DOCX = ROOT / "docs" / "okb_mvp_weekly_checklist_ru.docx"

PRIMARY = "1F4E79"
LIGHT = "EDF4FB"
LIGHT_GREEN = "EAF6F1"
LIGHT_GRAY = "F4F6F8"
TEXT = RGBColor(34, 34, 34)


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.color.rgb = TEXT

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Aptos"
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    list_bullet.font.size = Pt(11)
    list_bullet.font.color.rgb = TEXT
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(0)
    list_bullet.paragraph_format.line_spacing = 1.0


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_inline_runs(paragraph, text: str, size: int = 11) -> None:
    parts = text.split("`")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(size)
        if index % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.bold = True
        else:
            run.font.name = "Aptos"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")


def add_title(document: Document, date_line: str | None) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OKB MVP\nПонедельный план работ")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Практический план на 12 недель: сначала извлечение и распознавание документов, "
        "потом индексация, поиск и ответы со ссылками на источник."
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(88, 96, 105)
    subtitle.paragraph_format.space_after = Pt(8)

    if date_line:
        date_par = document.add_paragraph()
        date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_par.add_run(date_line)
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(110, 118, 129)
        date_par.paragraph_format.space_after = Pt(12)


def add_summary_box(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading = paragraph.add_run("Что внутри\n")
    heading.bold = True
    heading.font.size = Pt(12)
    heading.font.color.rgb = RGBColor(31, 78, 121)
    body = paragraph.add_run(
        "Два этапа MVP. Этап 1: импорт, маршрутизация файлов, прямое извлечение текста, "
        "распознавание текста, страницы, области страницы и атрибуты документа. "
        "Этап 2: фрагменты текста, индексация, поиск, ответы с источниками и осторожная сверка."
    )
    body.font.size = Pt(11)
    paragraph.paragraph_format.space_after = Pt(0)

    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_week_banner(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, PRIMARY)
    set_cell_margins(cell, top=90, start=140, bottom=90, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    paragraph.paragraph_format.space_after = Pt(0)


def add_section_label(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    fill = LIGHT_GREEN if "Готово" in text else LIGHT_GRAY
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=55, start=100, bottom=55, end=100)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(23, 53, 83) if "Готово" not in text else RGBColor(47, 125, 107)
    paragraph.paragraph_format.space_after = Pt(0)


def add_check_item(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_after = Pt(1)
    box = paragraph.add_run("☐ ")
    box.font.size = Pt(11)
    box.font.color.rgb = RGBColor(47, 125, 107)
    add_inline_runs(paragraph, text, size=11)


def add_bullet_item(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    add_inline_runs(paragraph, text, size=11)


def parse_markdown(md_text: str) -> tuple[str | None, list[tuple[str, str]]]:
    lines = md_text.splitlines()
    date_line = None
    items: list[tuple[str, str]] = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            items.append(("blank", ""))
            continue
        if stripped == "---":
            items.append(("rule", ""))
            continue
        if stripped.startswith("Дата:"):
            date_line = stripped
            continue
        if stripped.startswith("# "):
            items.append(("h1", stripped[2:].strip()))
            continue
        if stripped.startswith("## "):
            items.append(("h2", stripped[3:].strip()))
            continue
        if stripped.startswith("### "):
            items.append(("h3", stripped[4:].strip()))
            continue
        if stripped.startswith("- [ ] "):
            items.append(("check", stripped[6:].strip()))
            continue
        if stripped.startswith("- "):
            items.append(("bullet", stripped[2:].strip()))
            continue
        items.append(("text", stripped))

    return date_line, items


def build_document() -> Document:
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    date_line, items = parse_markdown(markdown)

    document = Document()
    set_document_defaults(document)
    add_title(document, date_line)
    add_summary_box(document)

    first_week = True
    started_main_content = False

    for kind, text in items:
        if kind == "h1":
            continue

        if kind == "h2":
            if text.startswith("Неделя "):
                if not first_week:
                    document.add_page_break()
                add_week_banner(document, text)
                first_week = False
                started_main_content = True
                spacer = document.add_paragraph()
                spacer.paragraph_format.space_after = Pt(1)
            else:
                heading = document.add_paragraph()
                heading_run = heading.add_run(text)
                heading_run.bold = True
                heading_run.font.size = Pt(14)
                heading_run.font.color.rgb = RGBColor(23, 53, 83)
                heading.paragraph_format.space_before = Pt(4 if started_main_content else 0)
                heading.paragraph_format.space_after = Pt(3)
            continue

        if kind == "h3":
            add_section_label(document, text)
            continue

        if kind == "check":
            add_check_item(document, text)
            continue

        if kind == "bullet":
            add_bullet_item(document, text)
            continue

        if kind == "text":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, text, size=11)
            continue

        if kind == "blank":
            continue

        if kind == "rule":
            document.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

    return document


def main() -> None:
    document = build_document()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
