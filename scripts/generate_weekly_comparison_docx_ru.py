# -*- coding: utf-8 -*-

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "okb_mvp_12_week_reference_comparison_ru.docx"


REFERENCE_LINKS = [
    ("DNV RuleAgent (ИИ-ассистент по правилам DNV)", "https://www.dnv.com/services/ruleagent/"),
    (
        "DNV Rules and Standards Explorer (навигатор по правилам и стандартам DNV)",
        "https://www.dnv.cn/rules-standards/",
    ),
    (
        "Bureau Veritas Rules Search (поиск по правилам Bureau Veritas)",
        "https://marine-offshore.bureauveritas.com/rules-search",
    ),
    (
        "Bureau Veritas Rules & Guidelines (правила и руководства Bureau Veritas)",
        "https://marine-offshore.bureauveritas.com/rules-guidelines",
    ),
    (
        "ABS Rule Manager 2.0 Quick Start Guide (краткое руководство по менеджеру правил ABS)",
        "https://ww2.eagle.org/content/dam/eagle/rules-and-resources/RuleManager2/rule-manager-2-quick-start-guide.pdf",
    ),
    ("OneOcean Regs4ships (нормативная библиотека для флота)", "https://www.oneocean.com/how-we-help/grc/regs4ships"),
    (
        "OneOcean MarineRegulations brochure (брошюра по MarineRegulations)",
        "https://www.oneocean.com/wp-content/uploads/2025/09/OneOcean_MarineRegulations_Sept2025-1.pdf",
    ),
    (
        "Azure Document Intelligence (сервис OCR и понимания документов)",
        "https://learn.microsoft.com/en-us/azure/ai-services/document-intelligence/overview",
    ),
    (
        "Azure AI Search agentic retrieval (агентный многошаговый поиск)",
        "https://learn.microsoft.com/en-us/azure/search/agentic-retrieval-overview",
    ),
    (
        "Azure OpenAI On Your Data (classic) (чаты по своим данным; deprecated на 20.04.2026)",
        "https://learn.microsoft.com/en-us/azure/foundry-classic/openai/concepts/use-your-data",
    ),
    (
        "Google Vertex AI Search parse and chunk documents (парсинг и разбиение документов)",
        "https://docs.cloud.google.com/generative-ai-app-builder/docs/parse-chunk-documents",
    ),
    (
        "Google Document AI extraction overview (обзор извлечения данных)",
        "https://docs.cloud.google.com/document-ai/docs/extracting-overview",
    ),
    (
        "OneOcean article on AI and regulatory content (статья об ИИ и регуляторном контенте)",
        "https://www.oneocean.com/insights/innovation-experts-on-fleet-futureproofing",
    ),
]


TERM_GLOSSARY = [
    ("MVP", "минимально жизнеспособный продукт"),
    ("Ingestion", "загрузка и первичная обработка"),
    ("OCR", "оптическое распознавание текста"),
    ("Layout parsing", "анализ структуры и макета документа"),
    ("Chunking", "разбиение документа на фрагменты"),
    ("Retrieval", "поиск релевантных фрагментов"),
    ("Hybrid retrieval / hybrid search", "гибридный поиск: точный плюс семантический"),
    ("RAG", "генерация ответа с опорой на найденные источники"),
    ("Citation", "цитата / ссылка на основание"),
    ("Reranking", "повторное ранжирование результатов"),
    ("Source registry", "реестр источников"),
    ("Frontend", "пользовательский интерфейс"),
    ("Benchmark set", "контрольный или эталонный набор запросов"),
    ("Evidence path", "цепочка оснований и проверяемых ссылок"),
]


REFERENCE_PRACTICES = [
    {
        "title": (
            "DNV RuleAgent (ИИ-ассистент по правилам DNV) и "
            "Rules and Standards Explorer (навигатор по правилам и стандартам)"
        ),
        "public_details": [
            "AI assistant (ИИ-ассистент) работает поверх официальной базы правил DNV, а не поверх случайной папки PDF.",
            "Публично раскрыто, что RuleAgent имеет direct integration with DNV rule databases (прямую интеграцию с базами правил DNV).",
            "Публично раскрыто, что ответы ведут к edition-specific references (ссылкам на конкретные редакции документов).",
            "Есть vessel-tailored guidance (подстройка под конкретное судно): используется информация о типе судна и class notations (классовых обозначениях).",
            "Публично раскрыто, что RuleAgent покрывает Rules for Ships и Rules for Offshore Units, а также связанные документы: CG, CP, SI, ST, OS и RP.",
        ],
        "not_public": [
            "Не раскрыты публично OCR engine (движок OCR), embedding model (модель эмбеддингов), vector database (векторная БД), схема chunking и конкретный reranker.",
            "Для OKB из этого решения лучше забирать продуктовую дисциплину: authoritative rule base (авторитетная база правил), навигацию по редакциям и привязку ответа к официальному источнику, а не пытаться угадывать их внутренний стек.",
        ],
        "sources": [
            ("RuleAgent", "https://www.dnv.com/services/ruleagent/"),
            ("Rules and Standards Explorer", "https://www.dnv.cn/rules-standards/"),
        ],
    },
    {
        "title": (
            "Bureau Veritas Rules Search (поиск по правилам) / "
            "Rules Explorer (навигатор по правилам)"
        ),
        "public_details": [
            "Публично виден упор на rules catalog (каталог правил) и structured search (структурированный поиск), а не на чат-интерфейс.",
            "Rules Search поддерживает поиск по Ref (коду документа), Reference (наименованию) и Edition date (дате редакции).",
            "Rules & Guidelines показывают latest updates (последние обновления) и дают доступ к большому набору актуальных технических публикаций.",
            "С точки зрения практики это сильный пример version-aware navigation (навигации с учётом версии и редакции документа).",
        ],
        "not_public": [
            "Публично не раскрыты OCR-стек, внутренний поисковый движок, схема chunking, embeddings или LLM-компоненты.",
            "На публичных страницах это скорее rule search platform (платформа поиска по правилам), чем раскрытая RAG-система.",
        ],
        "sources": [
            ("Rules Search", "https://marine-offshore.bureauveritas.com/rules-search"),
            ("Rules & Guidelines", "https://marine-offshore.bureauveritas.com/rules-guidelines"),
        ],
    },
    {
        "title": "ABS Rule Manager 2.0 (менеджер правил ABS)",
        "public_details": [
            "Публично раскрыто, что ABS Rule Manager gives fast and easily searchable access (даёт быстрый и удобно ищущийся доступ) к latest requirements (актуальным требованиям).",
            "Отдельно указано, что пользователи могут search the entire contents of the Rules (искать по полному содержимому правил) и browse publications based on those terms (просматривать публикации по найденным терминам).",
            "Это хороший ориентир по UX rules library (интерфейсу библиотеки правил): сначала быстрое нахождение требований, потом переход к документу.",
        ],
        "not_public": [
            "Публично не раскрыты OCR-движок, LLM, vector search, reranking или citation pipeline.",
            "По публичным материалам это не явная RAG-система, а продвинутый searchable rules library (поисковый каталог правил).",
        ],
        "sources": [
            (
                "Rule Manager 2.0 Quick Start Guide",
                "https://ww2.eagle.org/content/dam/eagle/rules-and-resources/RuleManager2/rule-manager-2-quick-start-guide.pdf",
            ),
            ("ABS Rules and Guides", "https://ww2.eagle.org/en/rules-and-resources/rules-and-guides-v2.html"),
        ],
    },
    {
        "title": (
            "OneOcean Regs4ships (нормативная библиотека для флота) / "
            "MarineRegulations (подборка применимых правил)"
        ),
        "public_details": [
            "Публично раскрыто, что Regs4ships хранит 100,000+ maritime documents (более 100 000 морских нормативных документов).",
            "Публично раскрыто, что updates are processed within 48 hours (обновления появляются в течение 48 часов после появления у источника).",
            "Есть verified, version-controlled documents (проверенные и версионируемые документы), online and full offline modes (онлайн и полноценный офлайн-режимы) и Compliance Index (контур обязательных документов для проверки/инспекции).",
            "MarineRegulations и Regs4ships строят tailored libraries (настраиваемые библиотеки) под flag, class, equipment, fuel, routes и trading area.",
            "В официальной статье OneOcean отдельно сказано, что в Regs4ships используется a set of Large Language Models (набор больших языковых моделей) для сопоставления curated, analyst-verified regulatory content (кураторски проверенного регуляторного контента) с реальными задачами моряков.",
        ],
        "not_public": [
            "Публично не раскрыты конкретный OCR engine, embedding model, vector database, схема индексации и точный RAG pipeline.",
            "Снаружи лучше всего видна не внутренняя библиотека компонентов, а продуктовая логика: curated source of truth (кураторский единый источник истины), быстрые обновления, офлайн-доступ и релевантные подборки документов.",
        ],
        "sources": [
            ("Regs4ships", "https://www.oneocean.com/how-we-help/grc/regs4ships"),
            ("MarineRegulations", "https://www.oneocean.com/how-we-help/grc/marineregulations"),
            (
                "OneOcean on AI in Regs4ships",
                "https://www.oneocean.com/insights/innovation-experts-on-fleet-futureproofing",
            ),
        ],
    },
    {
        "title": (
            "Azure AI Search + Document Intelligence + Azure OpenAI On Your Data "
            "(по состоянию на 20.04.2026 продукт в статусе deprecated)"
        ),
        "public_details": [
            "Здесь стек раскрыт лучше всего. Azure Document Intelligence имеет prebuilt-read (модель чтения), prebuilt-layout (модель анализа структуры), optional features `ocr.highResolution`, `ocr.formula`, `ocr.font`, `ocr.barcode`, поддержку keyValuePairs и searchable PDF.",
            "В Azure AI Search публично раскрыт integrated vectorization pipeline (встроенный конвейер векторизации): indexer pipeline (конвейер индексатора) -> Text Split skill (навык разбиения текста), Document Layout skill (навык анализа структуры документа) или document parsing modes (режимы парсинга документов) -> AzureOpenAIEmbeddingSkill / Vision VectorizeSkill -> vector fields (векторные поля индекса).",
            "Text Split skill умеет делить текст на sentences (предложения) или pages (страницы/фрагменты), выдаёт offsets (смещения), lengths (длины) и ordinal positions (порядковые позиции), поддерживает chunking по `azureOpenAITokens` и прямо показывает пример с `maximumPageLength: 512`.",
            "В документации Text Split skill отдельно указано, что skill реализует tokenization через tiktoken-подход с использованием SharpToken и `Microsoft.ML.Tokenizers`, а в примерах используется tokenizer `cl100k_base`.",
            "Azure OpenAI On Your Data публично раскрывает ingest flow: crack -> chunk -> embed -> index, field mapping (сопоставление полей), document-level access control (доступ на уровне документа), а также то, что для наилучшего результата обычно используется semantic search или hybrid + semantic search.",
            "Agentic retrieval (агентный/многошаговый поиск) официально описан так: LLM разбивает сложный запрос на subqueries (подзапросы), запускает их параллельно, затем применяется semantic reranking (семантическое повторное ранжирование), после чего результаты объединяются в единый ответ с source references (ссылками на источники).",
            "Важно по дате: в официальной документации Microsoft на 20 апреля 2026 года Azure OpenAI On Your Data (classic) помечен как deprecated and approaching retirement; Microsoft рекомендует миграцию в Foundry Agent Service с Foundry IQ. Для OKB это значит, что брать нужно архитектурные практики, а не копировать продукт буквально.",
        ],
        "not_public": [
            "Даже у Azure публично не раскрыты все внутренние параметры ранжирования и эвристики, но это самый прозрачный из всех рассмотренных стеков.",
        ],
        "sources": [
            ("Document Intelligence overview", "https://learn.microsoft.com/en-us/azure/ai-services/document-intelligence/overview"),
            ("Integrated vectorization", "https://learn.microsoft.com/en-us/azure/search/vector-search-integrated-vectorization"),
            ("Text Split skill", "https://learn.microsoft.com/en-us/azure/search/cognitive-search-skill-textsplit"),
            ("Integrated vectorization using REST", "https://learn.microsoft.com/en-us/azure/search/search-how-to-integrated-vectorization"),
            ("Azure OpenAI use your data", "https://learn.microsoft.com/en-us/azure/foundry-classic/openai/concepts/use-your-data"),
            ("Agentic retrieval", "https://learn.microsoft.com/en-us/azure/search/agentic-retrieval-overview"),
        ],
    },
    {
        "title": (
            "Google Document AI + Vertex AI Search "
            "(облачный референс по OCR, layout parsing и RAG)"
        ),
        "public_details": [
            "У Google публично раскрыты parser modes (режимы парсинга): digital parser (только машиночитаемый текст), OCR parser for PDFs (OCR-парсер для PDF) и layout parser (парсер структуры документа). Для OCR parser есть настройка `useNativeText`, которая позволяет смешивать OCR с уже доступным машинным текстом в PDF.",
            "OCR parser for PDFs рекомендуется для scanned PDFs и изображений с текстом, но Google отдельно пишет, что при сложной иерархии, таблицах и инфографике лучше использовать layout parser.",
            "Layout parser рекомендуется для rich content (богато структурированных документов), умеет обнаруживать paragraphs, tables, lists, titles, headings, headers и footnotes, а при включённом document chunking создаёт context-aware chunks (контекстно связные фрагменты) для retrieval и answer generation.",
            "В Vertex AI Search публично раскрыты layout-aware document chunking, `ChunkingConfig.LayoutBasedChunkingConfig`, `ParsingConfig.LayoutParsingConfig`, а также bring your own parsed document / bring your own chunks (загрузка заранее распарсенных документов или заранее подготовленных фрагментов).",
            "У layout parser есть add-ons: image annotation, table annotation и Gemini layout parsing. Для Gemini layout parsing публично заявлены better table recognition (лучшее распознавание таблиц), improved reading order (лучший порядок чтения) и more accurate text recognition (более точное распознавание текста).",
            "Google Document AI Form Parser публично раскрыт как компонент, который извлекает key-value pairs (ключ-значение), tables (таблицы) и selection marks / checkboxes (чекбоксы).",
            "Document AI Warehouse и Vertex AI Search публично раскрывают grounded summaries with citations (заземлённые сводки с цитатами), ACL/access control (контроль доступа) и confidence scores (оценки уверенности).",
        ],
        "not_public": [
            "Как и у Azure, не раскрываются все внутренние ранжирующие эвристики, но технологический стек документирован существенно лучше, чем у отраслевых морских продуктов.",
        ],
        "sources": [
            ("Parse and chunk documents", "https://docs.cloud.google.com/generative-ai-app-builder/docs/parse-chunk-documents"),
            ("Document AI extraction overview", "https://docs.cloud.google.com/document-ai/docs/extracting-overview"),
            ("Document AI handle response", "https://docs.cloud.google.com/document-ai/docs/handle-response"),
            ("Document AI Warehouse blog", "https://cloud.google.com/blog/products/ai-machine-learning/mobilize-your-unstructured-data-with-generative-ai"),
        ],
    },
]


TERM_REPLACEMENTS = [
    ("Azure OpenAI On Your Data (classic)", "Azure OpenAI On Your Data (classic, чаты по своим данным)"),
    ("Document Intelligence", "Document Intelligence (понимание документов)"),
    ("Rules and Standards Explorer", "Rules and Standards Explorer (навигатор по правилам и стандартам)"),
    ("RuleAgent", "RuleAgent (ИИ-ассистент по правилам)"),
    ("Regs4ships", "Regs4ships (нормативная библиотека для флота)"),
    ("MarineRegulations", "MarineRegulations (подборка применимых правил)"),
    ("agentic retrieval", "agentic retrieval (агентный многошаговый поиск)"),
    ("integrated vectorization", "integrated vectorization (встроенная векторизация)"),
    ("hybrid retrieval", "hybrid retrieval (гибридный поиск)"),
    ("hybrid search", "hybrid search (гибридный поиск)"),
    ("citation coverage", "citation coverage (полнота цитирования)"),
    ("source traceability", "source traceability (трассируемость до источника)"),
    ("benchmark set", "benchmark set (эталонный набор)"),
    ("page-level status", "page-level status (статус страницы)"),
    ("page-level confidence", "page-level confidence (уверенность на уровне страницы)"),
    ("text-first extraction", "text-first extraction (извлечение из текстового слоя в первую очередь)"),
    ("cited retrieval", "cited retrieval (поиск с обязательной привязкой к источникам)"),
    ("source constraints", "source constraints (ограничения по источникам)"),
    ("source adapters", "source adapters (адаптеры источников)"),
    ("search/chat", "search/chat (поиск и диалог)"),
    ("source preview", "source preview (предпросмотр источника)"),
    ("comparison result view", "comparison result view (экран результата сверки)"),
    ("review-only", "review-only (только для ревью инженером)"),
    ("evidence path", "evidence path (цепочка оснований)"),
    ("demo corpus", "demo corpus (демонстрационный корпус)"),
    ("reprocessing", "reprocessing (повторная обработка)"),
    ("chunking", "chunking (разбиение на фрагменты)"),
    ("chunk", "chunk (фрагмент)"),
    ("chunks", "chunks (фрагменты)"),
    ("retrieval", "retrieval (поиск релевантных фрагментов)"),
    ("lexical", "lexical (точный/лексический)"),
    ("metadata", "metadata (метаданные)"),
    ("semantic", "semantic (семантический)"),
    ("Frontend", "Frontend (пользовательский интерфейс)"),
    ("ingest pipeline", "ingest pipeline (конвейер загрузки и первичной обработки)"),
    ("query flow", "query flow (сценарий прохождения запроса)"),
    ("focused subqueries", "focused subqueries (узконаправленные подзапросы)"),
    ("page-aware", "page-aware (с учётом границ страницы)"),
    ("table-aware", "table-aware (с учётом структуры таблиц)"),
    ("fixed-size", "fixed-size (фиксированного размера)"),
    ("top-k", "top-k (верхние k результатов)"),
    ("grounded", "grounded (заземлённый на источниках)"),
    ("low-text", "low-text (страницы с очень малым объёмом текстового слоя)"),
    ("OCR-route", "OCR-route (маршрут обработки через OCR)"),
    ("processing state", "processing state (состояние обработки)"),
    ("inspection-ready", "inspection-ready (готовый к инспекционной проверке)"),
    ("audit-ready", "audit-ready (готовый к аудиту)"),
    ("offline mode", "offline mode (офлайн-режим)"),
    ("compliance", "compliance (соответствие нормативным требованиям)"),
    ("MVP", "MVP (минимально жизнеспособный продукт)"),
]


PUBLIC_IMPLEMENTATION_SUMMARY = [
    (
        "OCR и текст",
        [
            "У отраслевых морских решений DNV, Bureau Veritas, ABS и OneOcean конкретные OCR-движки публично не раскрыты. Снаружи видно не стек, а управляемую нормативную базу и поведение продукта.",
            "У Azure в роли OCR и document understanding публично описан Azure Document Intelligence: `prebuilt-read`, `prebuilt-layout`, optional features `ocr.highResolution`, `ocr.formula`, `ocr.font`, `ocr.barcode`, а также searchable PDF.",
            "У Google публично описаны `OCR parser for PDFs`, `Enterprise Document OCR`, `Layout parser` и режим `useNativeText`, когда OCR объединяется с уже имеющимся текстовым слоем PDF.",
        ],
    ),
    (
        "Таблицы и структура документа",
        [
            "Azure делает ставку на layout-first extraction (извлечение с учётом структуры): text, tables, paragraphs, key-value pairs и document structure извлекаются как отдельные сущности.",
            "Google Layout Parser и Form Parser публично описаны как инструменты для paragraphs, tables, lists, headings, page headers, footers, key-value pairs и selection marks.",
            "Vertex AI Search дополнительно раскрывает table annotation и Gemini layout parsing, где отдельно заявлены improved reading order и better table recognition.",
        ],
    ),
    (
        "Chunking и RAG",
        [
            "Azure раскрывает integrated vectorization, Text Split skill, Document Layout skill, AzureOpenAIEmbeddingSkill и agentic retrieval. В Text Split skill публично видны token chunking, overlap, offsets, lengths, ordinal positions и пример с `maximumPageLength: 512`.",
            "Google Vertex AI Search раскрывает layout-aware chunking, `ChunkingConfig.LayoutBasedChunkingConfig`, `ParsingConfig.LayoutParsingConfig` и режим bring your own chunks.",
            "У морских отраслевых продуктов RAG-схема публично не раскрыта на уровне библиотек и индексов, зато хорошо видны прикладные паттерны: version-aware navigation, authoritative source, tailored library, fast updates и обязательная привязка к официальной нормативной базе.",
        ],
    ),
]


WEEKS = [
    {
        "title": "Неделя 1. Подготовка",
        "baseline": (
            "Фиксируем границы MVP, инвентаризируем датасеты, собираем модель реестра "
            "источников и выбираем базовый стек OCR, парсинга, индексации и хранения."
        ),
        "references": (
            "DNV, Bureau Veritas и OneOcean начинают не с чата, а с управляемой нормативной "
            "базы: версия, источник, область применимости и статус редакции там являются "
            "частью модели данных. Azure-подход тоже требует заранее определить структуру индекса, "
            "поля для цитирования и модель доступа."
        ),
        "decision": (
            "В OKB уже на первой неделе принимаем как обязательные поля `source_authority`, "
            "`edition_date`, `authoritative_version`, `applicability_scope`, `document_status`, "
            "`processing_status`. Реестр документов считаем фундаментом, а не вспомогательной функцией."
        ),
    },
    {
        "title": "Неделя 2. Ingestion (загрузка и первичная обработка)",
        "baseline": (
            "Подключаем локальные папки и реестры, определяем типы файлов, запускаем text-first "
            "извлечение и формируем базовый ingest pipeline."
        ),
        "references": (
            "Azure и похожие enterprise-подходы рассматривают ingestion как документное понимание, "
            "а не просто вытаскивание текста. DNV и ABS, напротив, работают уже с хорошо "
            "подготовленным цифровым rule corpus. OneOcean делает акцент на контролируемом "
            "распространении актуальных документов."
        ),
        "decision": (
            "Оставляем в MVP text-first extraction плюс OCR fallback, но усиливаем ingest-метаданные: "
            "родословная источника, версия/редакция, кандидаты в дубликаты и page-level status."
        ),
    },
    {
        "title": "Неделя 3. OCR и page-level extraction (извлечение на уровне страницы)",
        "baseline": (
            "Обрабатываем low-text и scanned PDF, сохраняем трассировку до страницы, выделяем "
            "таблицы, блоки, ошибки и кандидатов на повторную обработку."
        ),
        "references": (
            "Azure Document Intelligence показывает, что для инженерных документов важны не только "
            "текст, но и layout, таблицы, key-value поля, штампы и структурные блоки. DNV и OneOcean "
            "по сути подтверждают ту же мысль с другой стороны: ценность возникает только тогда, "
            "когда извлечённый контент действительно пригоден для профессиональной проверки."
        ),
        "decision": (
            "Добавляем в OKB page-level confidence, low-text flags, признак OCR-route и структурные "
            "артефакты по таблицам и титульным блокам, где это возможно. Ошибочные страницы и файлы "
            "должны попадать в явную очередь reprocessing."
        ),
    },
    {
        "title": "Неделя 4. KB и chunking (база знаний и разбиение на фрагменты)",
        "baseline": (
            "Нормализуем текст, делим его на chunks, добавляем метаданные, строим embeddings и индекс."
        ),
        "references": (
            "Azure-стек опирается на корректно промапленные citation fields и структурное chunking. "
            "Bureau Veritas и ABS показывают, что поиск по точным кодам, разделам и редакциям "
            "критичен для нормативки. DNV сохраняет edition-specific navigation."
        ),
        "decision": (
            "Принимаем hybrid retrieval как базовую модель недели 4: lexical + metadata + semantic. "
            "Порядок chunking для OKB: раздел/пункт, затем page-aware chunks, затем table-aware "
            "фрагменты, и только в крайнем случае fallback fixed-size chunks."
        ),
    },
    {
        "title": "Неделя 5. Retrieval (поиск) и ответы с цитатами",
        "baseline": (
            "Запускаем cited retrieval, вводим source constraints и проверяем top-k качество на контрольном наборе."
        ),
        "references": (
            "DNV жёстко привязывает ответ к официальной rule base. Azure показывает, что качество "
            "цитирования зависит не только от промпта, но и от mapping полей, размера chunks и "
            "схемы индекса. OneOcean полезен как пример работы с версионностью и обновлениями."
        ),
        "decision": (
            "Для OKB закрепляем contract ответа: документ, страница, фрагмент/chunk и пометка о "
            "недостаточности основания, если evidence слабый. Размер chunks настраиваем по контрольным "
            "вопросам и качеству citation coverage, а не один раз 'на глаз'."
        ),
    },
    {
        "title": "Неделя 6. Диалог и уточнение контекста",
        "baseline": (
            "Делаем инженерный query flow: ассистент умеет уточнять слишком широкие вопросы и "
            "сохранять короткий контекст диалога."
        ),
        "references": (
            "Azure agentic retrieval умеет разбирать сложный запрос на несколько focused subqueries. "
            "DNV и OneOcean улучшают релевантность за счёт контекста судна, класса, операционной среды "
            "и применимого набора правил."
        ),
        "decision": (
            "В MVP оставляем диалог простым, но добавляем фильтры контекста: проект, тип судна, "
            "семейство правил, набор документов, период редакции. Multi-query retrieval включаем "
            "только для широких и составных вопросов."
        ),
    },
    {
        "title": "Неделя 7. Извлечение параметров для сверки",
        "baseline": (
            "Начинаем вытаскивать candidate parameters из чертежей и спецификаций и связывать их "
            "с нормативными требованиями."
        ),
        "references": (
            "Azure предлагает layout/field extraction и кастомные модели для сложных документов. "
            "DNV и OneOcean показывают, что точность сильно растёт, если retrieval знает доменный контекст. "
            "При этом ни один сильный референс не позиционирует ИИ как финального инженерного эксперта."
        ),
        "decision": (
            "Ограничиваем OKB в MVP малым списком формализуемых параметров: материал, толщина, масса, "
            "класс, обозначение и отдельные числовые величины. Каждый извлечённый параметр должен "
            "сохранять точную привязку к источнику."
        ),
    },
    {
        "title": "Неделя 8. Сверка требования и проектного значения",
        "baseline": (
            "Сравниваем проектные значения с нормой и классифицируем результат как `OK`, `WARNING` "
            "или `POTENTIAL_MISMATCH`."
        ),
        "references": (
            "DNV, Bureau Veritas и OneOcean сильны именно в трассируемости, версионности и professional "
            "verification. Общая лучшая практика в регулируемых областях: не выдавать автоматическое "
            "финальное одобрение."
        ),
        "decision": (
            "Сохраняем текущий принцип OKB: система не выносит окончательный инженерный вердикт. "
            "В интерфейсе обязательно показываем обе стороны сверки: нормативный фрагмент, проектный "
            "фрагмент, нормализованное значение/единицу и причину неоднозначности, если она есть."
        ),
    },
    {
        "title": "Неделя 9. Интеграции и source adapters (адаптеры источников)",
        "baseline": (
            "Формализуем импорты из локальных папок, таблиц-реестров, корпуса РКО и Google Drive; "
            "нормализуем метаданные и ищем дубликаты."
        ),
        "references": (
            "OneOcean силён в update propagation и раздаче релевантных нормативных наборов. "
            "Bureau Veritas показывает, как поиск живёт в связке с другими compliance-инструментами. "
            "Azure полезен как ориентир по ACL-aware retrieval и повторному использованию индексов."
        ),
        "decision": (
            "В MVP оставляем ограниченный набор интеграций, но усиливаем модель адаптера: `source_type`, "
            "`sync_mode`, `version_strategy`, `duplicate_keys`, поле под будущие ACL, а также задел "
            "под дорожку 'последние обновления'."
        ),
    },
    {
        "title": "Неделя 10. Frontend (пользовательский интерфейс) и рабочие экраны",
        "baseline": (
            "Собираем документный реестр, статус обработки, search/chat, cited answer view, source preview "
            "и comparison result view."
        ),
        "references": (
            "ABS показывает, что UX rules library сам по себе важен. Bureau Veritas даёт хороший пример "
            "поиска по reference/topic/edition. DNV усиливает линию source traceability, OneOcean — "
            "inspection-ready и audit-ready сценарии."
        ),
        "decision": (
            "Структурируем UI OKB вокруг трёх режимов: реестр, ответ с цитатами, сверка/ревью. "
            "Фильтры с первого дня: источник, код, редакция, проект и processing state."
        ),
    },
    {
        "title": "Неделя 11. Тестирование и оценка качества",
        "baseline": (
            "Проверяем top-k retrieval, citation coverage, OCR quality, latency и пользовательские сценарии."
        ),
        "references": (
            "Azure явно связывает качество поиска с дизайном индекса, chunking и mapping полей. "
            "OneOcean показывает ценность inspection readiness. DNV и ABS подчёркивают, что пользователю "
            "важна не только скорость ответа, но и доверие к evidence path."
        ),
        "decision": (
            "Делаем оценку видимой: контрольные запросы, доля top-3 релевантности, полнота цитирования, "
            "лог OCR-проблем и латентность. Добавляем небольшой экспертно-размеченный benchmark set, "
            "а не опираемся только на субъективное впечатление от демо."
        ),
    },
    {
        "title": "Неделя 12. Deploy (развёртывание) и повторяемое демо",
        "baseline": (
            "Доводим frontend, деплой, тесты и repeatable demo scenario до состояния воспроизводимого стенда."
        ),
        "references": (
            "OneOcean даёт полезный ориентир по offline/low-connectivity value, а Azure — по "
            "воспроизводимой архитектуре и модели доступа. DNV и Bureau Veritas показывают, что "
            "для доверия важнее стабильная навигация по авторитетным источникам, чем 'эффектный чат'."
        ),
        "decision": (
            "Оставляем деплой OKB простым и воспроизводимым. Замораживаем demo corpus, документируем "
            "import commands и проверяем repeatable cited answers. Полноценный offline mode можно "
            "отложить, но preview источника, видимость версии и видимость ошибок должны быть в демо."
        ),
    },
]


ADOPTED_NOW = [
    "Сначала реестр и версионность документов, потом ассистент.",
    "Hybrid retrieval: точный поиск по кодам и семантический поиск должны работать вместе.",
    "Structure-aware chunking: пункт, раздел, страница, таблица важнее, чем механическое деление по длине.",
    "Обязательный citation contract: документ, страница, фрагмент и пометка о слабом основании.",
    "OCR quality metadata: low-text, OCR-route, confidence и extraction failures должны быть видимы.",
    "Сверка только как engineer-review aid, без финального инженерного одобрения системой.",
    "Сужение retrieval по проекту, типу судна, семейству правил и источнику.",
]


DEFERRED = [
    "Document-level ACL в retrieval, если первый MVP живёт на доверенном внутреннем корпусе.",
    "Автоматическая доставка всех обновлений нормативки из источников.",
    "Полноценные динамические библиотеки под конкретное судно или проект.",
    "Offline library mode уровня OneOcean.",
    "Постоянный multi-query planning для каждого запроса.",
    "Глубокие кастомные extraction-модели для множества типов чертежей и спецификаций.",
]


OUT_OF_SCOPE = [
    "Нативное семантическое понимание DWG в MVP.",
    "Финальные инженерные заключения от системы.",
    "Позиционирование как общего enterprise-ассистента по всем данным подряд.",
    "Сложный интеграционный слой до стабилизации реестра и cited search.",
]


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def translate_terms(text: str) -> str:
    translated = text
    for source, target in sorted(TERM_REPLACEMENTS, key=lambda pair: len(pair[0]), reverse=True):
        translated = translated.replace(source, target)
    return translated


def add_label_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(translate_terms(text))
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(translate_terms(item))
        paragraph.paragraph_format.space_after = Pt(2)


def add_glossary(document: Document) -> None:
    for english, russian in TERM_GLOSSARY:
        paragraph = document.add_paragraph()
        term_run = paragraph.add_run(f"{english} ")
        term_run.bold = True
        paragraph.add_run(f"({russian})")
        paragraph.paragraph_format.space_after = Pt(3)


def add_reference_sources(document: Document, sources: list[tuple[str, str]]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Источники: ")
    run.bold = True
    for index, (name, url) in enumerate(sources):
        if index:
            paragraph.add_run("; ")
        add_hyperlink(paragraph, name, url)
    paragraph.paragraph_format.space_after = Pt(6)


def add_reference_practices(document: Document) -> None:
    for practice in REFERENCE_PRACTICES:
        document.add_heading(practice["title"], level=2)
        add_label_paragraph(
            document,
            "Что публично подтверждено",
            (
                "Ниже перечислены только те детали, которые прямо видны в официальных описаниях, "
                "публичной документации или официальных брошюрах."
            ),
        )
        add_bullets(document, practice["public_details"])
        add_label_paragraph(
            document,
            "Чего публично не видно",
            (
                "Если вендор не раскрыл точный OCR-движок, embeddings, vector database или схему "
                "RAG, в документе это помечено явно и без домыслов."
            ),
        )
        add_bullets(document, practice["not_public"])
        add_reference_sources(document, practice["sources"])


def add_public_implementation_summary(document: Document) -> None:
    for title, bullets in PUBLIC_IMPLEMENTATION_SUMMARY:
        document.add_heading(title, level=2)
        add_bullets(document, bullets)


def build_document() -> Document:
    document = Document()
    set_document_defaults(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OKB MVP: сравнительный 12-недельный план\nс учётом лучших внешних решений")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Рабочая версия на русском языке. Дата: 20.04.2026")
    subtitle.paragraph_format.space_after = Pt(12)

    document.add_heading("Назначение документа", level=1)
    document.add_paragraph(
        "Этот документ нужен для управленческого сравнения текущего плана OKB с лучшими "
        "публичными решениями, которые ближе всего к нашему контуру: инженерная и нормативная "
        "документация, OCR (оптическое распознавание текста), knowledge base (база знаний), "
        "ответы с обязательными цитатами, контроль версий "
        "и ограниченная сверка параметров без автоматического инженерного вердикта."
    )
    document.add_paragraph(
        "По каждой неделе ниже показаны три вещи: что уже заложено в план OKB, что делают сильные "
        "внешние решения, и что именно стоит забирать в наш MVP."
    )
    document.add_paragraph(
        "Рядом с английскими названиями и ключевыми техническими терминами по возможности дано "
        "русское пояснение. Если точное внутреннее устройство внешнего решения публично не раскрыто, "
        "это указано отдельно."
    )

    document.add_heading("Использованные внешние ориентиры", level=1)
    for name, url in REFERENCE_LINKS:
        paragraph = document.add_paragraph(style="List Bullet")
        add_hyperlink(paragraph, name, url)

    document.add_heading("Как читать внешние практики", level=1)
    document.add_paragraph(
        "У отраслевых решений DNV, Bureau Veritas, ABS и OneOcean публично обычно раскрывается не "
        "внутренний ML-стек, а то, как продукт работает для пользователя: где находится официальный "
        "источник правил, как учитывается редакция документа, как выглядит навигация и как обеспечивается "
        "доверие к результату."
    )
    document.add_paragraph(
        "У облачных стеков Azure и Google официальная документация раскрывает конкретные управляемые "
        "компоненты гораздо подробнее: OCR-модели, layout parsing, chunking-конфигурации, embedding skills, "
        "field mapping, доступ на уровне документа и многошаговый retrieval. Поэтому именно их стоит "
        "использовать как референс по технической архитектуре."
    )
    document.add_paragraph(
        "Важно по состоянию на 20.04.2026: Azure OpenAI On Your Data (чаты по своим данным) "
        "в официальной документации Microsoft уже помечен как deprecated and approaching retirement. "
        "Для OKB полезны его архитектурные паттерны RAG, но не буквальное копирование продукта."
    )

    document.add_heading("Краткий глоссарий терминов", level=1)
    add_glossary(document)

    document.add_heading("Что по техреализации видно лучше всего", level=1)
    add_public_implementation_summary(document)

    document.add_heading("Что публично известно о технической реализации внешних решений", level=1)
    add_reference_practices(document)

    document.add_heading("Сравнение по неделям", level=1)
    for week in WEEKS:
        document.add_heading(week["title"], level=2)
        add_label_paragraph(document, "Что уже принято в OKB", week["baseline"])
        add_label_paragraph(document, "Что делают сильные внешние решения", week["references"])
        add_label_paragraph(document, "Что берём в OKB", week["decision"])

    document.add_heading("Что принимаем в основу уже сейчас", level=1)
    add_bullets(document, ADOPTED_NOW)

    document.add_heading("Что можно отложить после MVP", level=1)
    add_bullets(document, DEFERRED)

    document.add_heading("Что сознательно не обещаем в MVP", level=1)
    add_bullets(document, OUT_OF_SCOPE)

    document.add_heading("Короткий вывод", level=1)
    document.add_paragraph(
        "Главный вывод из сравнения такой: OKB уже движется в правильную сторону, но лучшее, "
        "что можно забрать из внешних решений, связано не с 'магией ИИ', а с дисциплиной корпуса. "
        "Нужно раньше и жёстче закрепить управляемый реестр, версии, source traceability, "
        "hybrid retrieval, quality metadata по OCR и режим review-only для инженерной сверки."
    )
    document.add_paragraph(
        "Иными словами, сильный MVP OKB — это не просто чат по PDF, а управляемая инженерная "
        "knowledge base с проверяемыми основаниями."
    )

    return document


def main() -> None:
    document = build_document()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
