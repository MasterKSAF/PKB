# -*- coding: utf-8 -*-
"""Extract readable text and build a knowledge-base reading report."""

from __future__ import annotations

import csv
import hashlib
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import docx
import fitz
from openpyxl import load_workbook
from striprtf.striprtf import rtf_to_text


ROOT = Path(__file__).resolve().parents[1]
DATA_DIRS = [ROOT / "data" / "gdrive_source", ROOT / "data" / "rko_izdaniya"]
DOCS_DIR = ROOT / "docs"
OUT_DIR = ROOT / "data" / "extracted_text"
REPORT_CSV = DOCS_DIR / "knowledge_base_reading_report.csv"
REPORT_MD = DOCS_DIR / "knowledge_base_reading_report.md"
LOW_TEXT_THRESHOLD_PER_PAGE = 100


def clean_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def safe_rel(path: Path) -> Path:
    try:
        return path.relative_to(ROOT)
    except ValueError:
        return path


def text_path_for(path: Path) -> Path:
    rel = safe_rel(path)
    digest = hashlib.sha1(str(rel).encode("utf-8")).hexdigest()[:10]
    name = re.sub(r"[^0-9A-Za-zа-яА-ЯёЁ._-]+", "_", path.stem).strip("_")[:120]
    source = classify_source(path)
    # Keep output paths short enough for Windows even when source folders are deeply nested.
    return OUT_DIR / source / f"{name}__{digest}.txt"


def write_text(path: Path, text: str) -> Path | None:
    if not text.strip():
        return None
    out = text_path_for(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8", errors="ignore")
    return out


def extract_pdf(path: Path) -> dict:
    doc = fitz.open(path)
    page_texts: list[str] = []
    page_chars: list[int] = []
    for page_index in range(doc.page_count):
        text = clean_text(doc.load_page(page_index).get_text())
        page_texts.append(f"\n\n===== PAGE {page_index + 1} =====\n{text}")
        page_chars.append(len(text))
    text = "".join(page_texts).strip()
    low_text_pages = sum(1 for chars in page_chars if chars < LOW_TEXT_THRESHOLD_PER_PAGE)
    avg_chars = sum(page_chars) / max(1, len(page_chars))
    status = "read"
    if len(text) == 0:
        status = "needs_ocr"
    elif low_text_pages == doc.page_count:
        status = "likely_scan_or_low_text"
    elif low_text_pages:
        status = "partial_low_text"
    out = write_text(path, text)
    result = {
        "pages": doc.page_count,
        "chars": len(text),
        "avg_chars_per_page": round(avg_chars, 1),
        "low_text_pages": low_text_pages,
        "status": status,
        "extracted_path": str(safe_rel(out)) if out else "",
        "sample": clean_text(text[:500]),
    }
    doc.close()
    return result


def extract_docx(path: Path) -> dict:
    document = docx.Document(path)
    parts: list[str] = []
    for para in document.paragraphs:
        text = clean_text(para.text)
        if text:
            parts.append(text)
    for table_no, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"\n===== TABLE {table_no} =====\n" + "\n".join(rows))
    text = "\n".join(parts)
    out = write_text(path, text)
    return {
        "pages": "",
        "chars": len(text),
        "avg_chars_per_page": "",
        "low_text_pages": "",
        "status": "read" if text else "empty",
        "extracted_path": str(safe_rel(out)) if out else "",
        "sample": clean_text(text[:500]),
    }


def extract_xlsx(path: Path) -> dict:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"\n===== SHEET {sheet.title} =====")
        for row in sheet.iter_rows(values_only=True):
            cells = [clean_text(str(value)) for value in row if value is not None and clean_text(str(value))]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    out = write_text(path, text)
    return {
        "pages": "",
        "chars": len(text),
        "avg_chars_per_page": "",
        "low_text_pages": "",
        "status": "read" if text else "empty",
        "extracted_path": str(safe_rel(out)) if out else "",
        "sample": clean_text(text[:500]),
    }


def extract_rtf(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    text = clean_text(rtf_to_text(raw))
    # Some RTF files contain mojibake after extraction. Keep the text, but flag it.
    status = "read_mojibake_possible" if "Ã" in text[:200] or "Ð" in text[:200] else "read"
    out = write_text(path, text)
    return {
        "pages": "",
        "chars": len(text),
        "avg_chars_per_page": "",
        "low_text_pages": "",
        "status": status if text else "empty",
        "extracted_path": str(safe_rel(out)) if out else "",
        "sample": clean_text(text[:500]),
    }


def extract_txt(path: Path) -> dict:
    text = clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    out = write_text(path, text)
    return {
        "pages": "",
        "chars": len(text),
        "avg_chars_per_page": "",
        "low_text_pages": "",
        "status": "read" if text else "empty",
        "extracted_path": str(safe_rel(out)) if out else "",
        "sample": clean_text(text[:500]),
    }


def classify_source(path: Path) -> str:
    rel = safe_rel(path).as_posix()
    if rel.startswith("data/gdrive_source/"):
        return "gdrive_source"
    if rel.startswith("data/rko_izdaniya/"):
        return "rko_izdaniya"
    return "other"


def process_file(path: Path) -> dict:
    ext = path.suffix.lower()
    base = {
        "source": classify_source(path),
        "path": str(safe_rel(path)),
        "name": path.name,
        "extension": ext,
        "size_bytes": path.stat().st_size,
        "size_mb": round(path.stat().st_size / 1024 / 1024, 3),
        "pages": "",
        "chars": "",
        "avg_chars_per_page": "",
        "low_text_pages": "",
        "status": "",
        "extracted_path": "",
        "sample": "",
        "error": "",
    }
    try:
        if ext == ".pdf":
            base.update(extract_pdf(path))
        elif ext == ".docx":
            base.update(extract_docx(path))
        elif ext == ".xlsx":
            base.update(extract_xlsx(path))
        elif ext == ".rtf":
            base.update(extract_rtf(path))
        elif ext == ".txt":
            base.update(extract_txt(path))
        elif ext == ".doc":
            base["status"] = "needs_doc_converter"
        elif ext == ".dwg":
            base["status"] = "needs_cad_converter"
        elif ext == ".csv":
            base.update(extract_txt(path))
        else:
            base["status"] = "unsupported"
    except Exception as exc:  # noqa: BLE001
        base["status"] = "error"
        base["error"] = str(exc)
    return base


def iter_files() -> list[Path]:
    files: list[Path] = []
    for directory in DATA_DIRS:
        if directory.exists():
            for path in directory.rglob("*"):
                if path.is_file() and ".part" not in path.name and "_meta" not in path.parts:
                    files.append(path)
    return sorted(files, key=lambda p: str(safe_rel(p)).lower())


def write_reports(records: list[dict]) -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "source",
        "path",
        "name",
        "extension",
        "size_bytes",
        "size_mb",
        "pages",
        "chars",
        "avg_chars_per_page",
        "low_text_pages",
        "status",
        "extracted_path",
        "sample",
        "error",
    ]
    with REPORT_CSV.open("w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(records)

    by_source = Counter(record["source"] for record in records)
    by_ext = Counter(record["extension"] for record in records)
    by_status = Counter(record["status"] for record in records)
    total_size = sum(record["size_bytes"] for record in records)
    total_chars = sum(int(record["chars"] or 0) for record in records)
    total_pages = sum(int(record["pages"] or 0) for record in records)

    low_text_pdfs = [
        record for record in records
        if record["extension"] == ".pdf" and record["status"] in {"needs_ocr", "likely_scan_or_low_text", "partial_low_text"}
    ]
    unsupported = [
        record for record in records
        if record["status"] in {"needs_doc_converter", "needs_cad_converter", "unsupported", "error"}
    ]
    largest = sorted(records, key=lambda item: item["size_bytes"], reverse=True)[:15]
    most_text = sorted(records, key=lambda item: int(item["chars"] or 0), reverse=True)[:15]

    lines = [
        "# Knowledge base reading report",
        "",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Summary",
        "",
        f"- Files processed: {len(records)}",
        f"- Total source size: {total_size / 1024 / 1024:.2f} MB",
        f"- Extracted text characters: {total_chars:,}".replace(",", " "),
        f"- PDF pages seen: {total_pages:,}".replace(",", " "),
        f"- Extracted text folder: `{safe_rel(OUT_DIR)}`",
        "",
        "## By source",
        "",
    ]
    for source, count in sorted(by_source.items()):
        lines.append(f"- `{source}`: {count}")
    lines.extend(["", "## By extension", ""])
    for ext, count in sorted(by_ext.items()):
        lines.append(f"- `{ext}`: {count}")
    lines.extend(["", "## By status", ""])
    for status, count in sorted(by_status.items()):
        lines.append(f"- `{status}`: {count}")

    lines.extend(["", "## Low-text / OCR candidates", ""])
    if low_text_pdfs:
        for record in low_text_pdfs:
            lines.append(
                f"- `{record['status']}`: [{record['name']}]({ROOT / record['path']}) "
                f"pages={record['pages']} chars={record['chars']} low_text_pages={record['low_text_pages']}"
            )
    else:
        lines.append("- None detected.")

    lines.extend(["", "## Unsupported or converter-required files", ""])
    if unsupported:
        for record in unsupported:
            lines.append(f"- `{record['status']}`: [{record['name']}]({ROOT / record['path']})")
    else:
        lines.append("- None.")

    lines.extend(["", "## Largest files", ""])
    for record in largest:
        lines.append(f"- {record['size_mb']} MB: [{record['name']}]({ROOT / record['path']})")

    lines.extend(["", "## Most text extracted", ""])
    for record in most_text:
        chars = int(record["chars"] or 0)
        if chars:
            lines.append(f"- {chars:,} chars: [{record['name']}]({ROOT / record['path']})".replace(",", " "))

    lines.extend(
        [
            "",
            "## Reading notes",
            "",
            "- Text-first extraction was performed for PDFs with embedded text, DOCX, XLSX, RTF, TXT, and CSV.",
            "- Low-text PDF files are readable only after OCR; they were detected but not OCR-processed in this pass.",
            "- DWG files require CAD conversion/parsing and were not semantically read.",
            "- Legacy DOC requires a DOC converter, such as LibreOffice or antiword, and was not fully read in this pass.",
        ]
    )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = iter_files()
    records = []
    for index, path in enumerate(files, start=1):
        print(f"[{index}/{len(files)}] {safe_rel(path)}")
        records.append(process_file(path))
    write_reports(records)
    print(json.dumps({
        "files": len(records),
        "report_csv": str(REPORT_CSV),
        "report_md": str(REPORT_MD),
        "out_dir": str(OUT_DIR),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
